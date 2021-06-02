import os
import time
from configparser import ConfigParser

import docx2pdf
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt , Inches
from docx.text import font


# import pythoncom


# creating the docx and PDF file with thr images downloaded and adding few details to report
# imagepath path where the images are downloaded
# docname-> name of the document to be created
# testlist->list tuple which contains the data to be added to report first page


def createreport(imagepath , docname , testlist) :
    print ( testlist )
    font.name = 'Calibri'
    font.size = Pt ( 12 )

    output = "output\\"
    imgDir = imagepath
    filenames = os.listdir ( imgDir )  #
    document = Document ( )
    paragraph = document.add_paragraph ( )

    # reportheading = "Performance report"
    file = 'config.ini'
    config = ConfigParser ( )
    config.read ( file )
    reportheading = (config [ 'report' ] [ 'heading' ])
    print ( reportheading )
    reportname = docname
    document.add_heading ( reportheading , 1 ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Creating a table object
    table = document.add_table ( rows=1 , cols=2 )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Adding heading in the 1st row of the table
    row = table.rows [ 0 ].cells
    row [ 0 ].text = ''
    row [ 1 ].text = ''

    # Adding data from the list to the table
    for id , name in testlist :
        # Adding a row and then adding data in it.
        row = table.add_row ( ).cells
        # Converting id to string as table can only take string input
        row [ 0 ].text = str ( id )
        row [ 1 ].text = name
    document.add_page_break ( )
    for index , filename in enumerate ( filenames ) :  # loop through all the files and folders for adding pictures
        document.add_picture ( imgDir + filename , width=Inches ( 7 ) )
        document.save ( output + reportname )

    time.sleep ( 2 )
    # creating the pdf from the  docx created
    docx2pdf.convert ( output + reportname )

    #doc2pdf =  output+reportname
    #print ( doc2pdf )
    #output = subprocess.check_output ( [ 'libreoffice' , '--convert-to' , 'pdf' , doc2pdf ],shell=True )
    #print ( output )
